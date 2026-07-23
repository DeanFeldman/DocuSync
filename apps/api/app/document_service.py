from __future__ import annotations

import hashlib
import math
import re
import shutil
import subprocess
import threading
import zipfile
from collections import defaultdict
from io import BytesIO
from pathlib import Path
from uuid import uuid4

from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph
from fastapi import HTTPException, UploadFile, status
from sqlalchemy import func, select
from sqlalchemy.orm import Session, selectinload
from datetime import datetime, timezone

from .config import settings
from .models import (
    DocumentElement,
    DocumentRecord,
    DocumentSet,
    GeneratedVersion,
    GenerationJob,
    GenerationTarget,
    LinkGroup,
    LinkMember,
)


SAFE_NAME_PATTERN = re.compile(r"[^A-Za-z0-9._ -]+")
WHITESPACE_PATTERN = re.compile(r"\s+")
TABLE_CELL_STYLE_PATTERN = re.compile(r"^table_cell:(\d+):(\d+):(\d+)$")
ORDERED_TABLE_CELL_STYLE_PATTERN = re.compile(
    r"^table_cell_order:(\d+):(\d+):(\d+):(\d+)$"
)
ORDERED_BODY_STYLE_PATTERN = re.compile(r"^body_order:(\d+):(.*)$", re.DOTALL)
PAGE_LAYOUT_UNITS = 18
WORD_RENDER_LOCK = threading.Lock()
LARGE_SET_LINK_GROUP_LIMIT = 100

def utc_isoformat(value: datetime) -> str:
    """Return a consistent timezone-aware UTC timestamp."""

    if value.tzinfo is None:
        value = value.replace(tzinfo=timezone.utc)
    else:
        value = value.astimezone(timezone.utc)

    return value.isoformat()

def new_id() -> str:
    return str(uuid4())


def normalise_text(text: str) -> str:
    return WHITESPACE_PATTERN.sub(" ", text).strip().casefold()


def table_cell_location(style_name: str | None) -> tuple[int, int, int] | None:
    value = style_name or ""
    ordered_match = ORDERED_TABLE_CELL_STYLE_PATTERN.fullmatch(value)
    if ordered_match is not None:
        _order, table_index, row_index, column_index = ordered_match.groups()
        return int(table_index), int(row_index), int(column_index)

    match = TABLE_CELL_STYLE_PATTERN.fullmatch(value)
    if match is None:
        return None
    return tuple(int(item) for item in match.groups())


def element_document_order(element: DocumentElement) -> tuple[int, int, int]:
    style_name = element.style_name or ""

    table_match = ORDERED_TABLE_CELL_STYLE_PATTERN.fullmatch(style_name)
    if table_match is not None:
        order, _table_index, row_index, column_index = table_match.groups()
        return int(order), int(row_index), int(column_index)

    body_match = ORDERED_BODY_STYLE_PATTERN.fullmatch(style_name)
    if body_match is not None:
        return int(body_match.group(1)), 0, 0

    return element.paragraph_index, 0, 0


def display_style_name(style_name: str | None) -> str | None:
    body_match = ORDERED_BODY_STYLE_PATTERN.fullmatch(style_name or "")
    if body_match is None:
        return style_name
    return body_match.group(2) or None


def element_location_payload(style_name: str | None) -> dict:
    location = table_cell_location(style_name)
    if location is None:
        return {}
    table_index, row_index, column_index = location
    return {
        "table_index": table_index,
        "row_index": row_index,
        "column_index": column_index,
    }


def element_type_for_style(style_name: str | None) -> str:
    if table_cell_location(style_name) is not None:
        return "table_cell"

    style = (display_style_name(style_name) or "").casefold()
    if style.startswith("heading") or style.startswith("title"):
        return "heading"
    if style.startswith("list"):
        return "list_item"
    return "paragraph"


def _layout_pages(elements: list[DocumentElement]) -> list[dict]:
    """Create honest, deterministic logical pages for the structured selection preview.

    python-docx does not calculate Word pagination. These page groups give the browser a
    stable scrolling surface without pretending to reproduce Word's layout engine.
    """

    pages: list[dict] = []
    current: list[dict] = []
    used_units = 0
    for element in sorted(elements, key=element_document_order):
        element_type = element_type_for_style(element.style_name)
        units = max(1, math.ceil(len(element.text) / 120))
        if element_type == "heading":
            units += 1
        if current and used_units + units > PAGE_LAYOUT_UNITS:
            pages.append({"page_number": len(pages) + 1, "elements": current})
            current = []
            used_units = 0
        current.append(
            {
                "id": element.id,
                "document_id": element.document_id,
                "paragraph_index": element.paragraph_index,
                "element_type": element_type,
                "text": element.text,
                "style_name": display_style_name(element.style_name),
                **element_location_payload(element.style_name),
                "page_number": len(pages) + 1,
            }
        )
        used_units += units
    if current or not pages:
        pages.append({"page_number": len(pages) + 1, "elements": current})
    return pages


def safe_download_name(name: str) -> str:
    cleaned = SAFE_NAME_PATTERN.sub("_", Path(name).name).strip(" .")
    return cleaned or "document.docx"


def _validate_docx_payload(filename: str, payload: bytes) -> None:
    if not filename.lower().endswith(".docx"):
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=f"{filename}: only DOCX files are supported.",
        )
    if len(payload) > settings.max_file_bytes:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"{filename}: file exceeds the configured size limit.",
        )
    if not zipfile.is_zipfile(BytesIO(payload)):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"{filename}: the file is not a valid DOCX archive.",
        )

    with zipfile.ZipFile(BytesIO(payload)) as archive:
        entries = archive.infolist()
        names = {entry.filename for entry in entries}
        if "[Content_Types].xml" not in names or "word/document.xml" not in names:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"{filename}: required DOCX parts are missing.",
            )
        if len(entries) > 1_000:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"{filename}: the DOCX archive contains too many parts.",
            )
        if any(entry.flag_bits & 0x1 for entry in entries):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"{filename}: password-protected DOCX files are not supported.",
            )
        total_uncompressed = sum(entry.file_size for entry in entries)
        if total_uncompressed > 100 * 1024 * 1024:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"{filename}: the expanded DOCX archive is too large.",
            )


def _extract_paragraphs(payload: bytes) -> list[tuple[int, str, str | None]]:
    try:
        document = Document(BytesIO(payload))
    except Exception as exc:  # python-docx raises several package/XML errors
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="A DOCX file could not be opened. It may be corrupt or unsupported.",
        ) from exc

    elements: list[tuple[int, str, str | None]] = []
    body_paragraph_index = 0
    table_index = 0
    document_order = 0
    synthetic_index = len(document.paragraphs)

    paragraph_lookup = {
        paragraph._p: paragraph
        for paragraph in document.paragraphs
    }
    table_lookup = {
        table._tbl: table
        for table in document.tables
    }

    for child in document.element.body.iterchildren():
        paragraph = paragraph_lookup.get(child)
        if paragraph is not None:
            text = paragraph.text.strip()
            if text:
                style_name = paragraph.style.name if paragraph.style is not None else ""
                elements.append(
                    (
                        body_paragraph_index,
                        text,
                        f"body_order:{document_order}:{style_name}",
                    )
                )
            body_paragraph_index += 1
            document_order += 1
            continue

        table = table_lookup.get(child)
        if table is None:
            continue

        seen_cells: set[object] = set()
        for row_index, row in enumerate(table.rows):
            for column_index, cell in enumerate(row.cells):
                cell_key = cell._tc
                if cell_key in seen_cells:
                    continue
                seen_cells.add(cell_key)

                text = cell.text.strip()
                if not text:
                    continue

                elements.append(
                    (
                        synthetic_index,
                        text,
                        (
                            f"table_cell_order:{document_order}:{table_index}:"
                            f"{row_index}:{column_index}"
                        ),
                    )
                )
                synthetic_index += 1

        table_index += 1
        document_order += 1

    return elements

async def create_document_set(
    session: Session,
    name: str,
    files: list[UploadFile],
) -> DocumentSet:
    cleaned_name = name.strip()
    if not cleaned_name:
        raise HTTPException(status_code=422, detail="Document-set name is required.")
    if len(files) < 2:
        raise HTTPException(status_code=422, detail="Upload at least two DOCX files.")
    if len(files) > settings.max_files_per_set:
        raise HTTPException(
            status_code=422,
            detail=f"A document set may contain at most {settings.max_files_per_set} files.",
        )

    document_set = DocumentSet(id=new_id(), name=cleaned_name)
    session.add(document_set)
    original_dir = settings.data_dir / "originals" / document_set.id
    original_dir.mkdir(parents=True, exist_ok=False)

    seen_names: set[str] = set()
    try:
        for upload in files:
            original_name = safe_download_name(upload.filename or "document.docx")
            name_key = original_name.casefold()
            if name_key in seen_names:
                raise HTTPException(
                    status_code=422,
                    detail=f"Duplicate file name in document set: {original_name}.",
                )
            seen_names.add(name_key)
            payload = await upload.read()
            _validate_docx_payload(original_name, payload)
            extracted = _extract_paragraphs(payload)

            document_id = new_id()
            stored_name = f"{document_set.id}/{document_id}.docx"
            target = settings.data_dir / "originals" / stored_name
            target.write_bytes(payload)

            record = DocumentRecord(
                id=document_id,
                document_set=document_set,
                original_name=original_name,
                stored_name=stored_name,
                checksum_sha256=hashlib.sha256(payload).hexdigest(),
            )
            session.add(record)

            for paragraph_index, text, style_name in extracted:
                session.add(
                    DocumentElement(
                        id=new_id(),
                        document=record,
                        paragraph_index=paragraph_index,
                        text=text,
                        normalized_text=normalise_text(text),
                        style_name=style_name,
                    )
                )

        session.flush()
        _create_exact_link_groups(session, document_set.id)
        session.commit()
    except Exception:
        session.rollback()
        shutil.rmtree(original_dir, ignore_errors=True)
        raise

    return get_document_set_or_404(session, document_set.id)


async def add_documents_to_set(
    session: Session,
    document_set_id: str,
    files: list[UploadFile],
) -> DocumentSet:
    # Add DOCX files to an existing set and rebuild exact-match groups.
    document_set = get_document_set_or_404(session, document_set_id)
    if not files:
        raise HTTPException(status_code=422, detail="Choose at least one DOCX file.")

    resulting_count = len(document_set.documents) + len(files)
    if resulting_count > settings.max_files_per_set:
        raise HTTPException(
            status_code=422,
            detail=(
                f"A document set may contain at most {settings.max_files_per_set} files. "
                f"This set currently contains {len(document_set.documents)}."
            ),
        )

    seen_names = {document.original_name.casefold() for document in document_set.documents}
    created_paths: list[Path] = []

    try:
        for upload in files:
            original_name = safe_download_name(upload.filename or "document.docx")
            name_key = original_name.casefold()
            if name_key in seen_names:
                raise HTTPException(
                    status_code=422,
                    detail=f"A document named {original_name} already exists in this set.",
                )
            seen_names.add(name_key)

            payload = await upload.read()
            _validate_docx_payload(original_name, payload)
            extracted = _extract_paragraphs(payload)

            document_id = new_id()
            stored_name = f"{document_set.id}/{document_id}.docx"
            target = settings.data_dir / "originals" / stored_name
            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_bytes(payload)
            created_paths.append(target)

            record = DocumentRecord(
                id=document_id,
                document_set=document_set,
                original_name=original_name,
                stored_name=stored_name,
                checksum_sha256=hashlib.sha256(payload).hexdigest(),
            )
            session.add(record)
            for paragraph_index, text, style_name in extracted:
                session.add(
                    DocumentElement(
                        id=new_id(),
                        document=record,
                        paragraph_index=paragraph_index,
                        text=text,
                        normalized_text=normalise_text(text),
                        style_name=style_name,
                    )
                )

        session.flush()
        _rebuild_exact_link_groups(session, document_set_id)
        session.commit()
        session.expire_all()
    except Exception:
        session.rollback()
        for path in created_paths:
            path.unlink(missing_ok=True)
        raise

    return get_document_set_or_404(session, document_set_id)


def remove_document_from_set(
    session: Session,
    document_set_id: str,
    document_id: str,
) -> DocumentSet:
    # Remove one document while preserving the minimum two-document set.
    document_set = get_document_set_or_404(session, document_set_id)
    if len(document_set.documents) <= 2:
        raise HTTPException(
            status_code=422,
            detail="A document set must keep at least two documents.",
        )

    document = next(
        (item for item in document_set.documents if item.id == document_id),
        None,
    )
    if document is None:
        raise HTTPException(status_code=404, detail="Document not found in this set.")

    original_path = settings.data_dir / "originals" / document.stored_name
    rendered_path = rendered_pdf_path(document)

    try:
        session.delete(document)
        session.flush()
        _rebuild_exact_link_groups(session, document_set_id)
        session.commit()
        session.expire_all()
    except Exception:
        session.rollback()
        raise

    original_path.unlink(missing_ok=True)
    rendered_path.unlink(missing_ok=True)
    return get_document_set_or_404(session, document_set_id)


def delete_document_set(session: Session, document_set_id: str) -> dict:
    # Permanently delete a local set and its stored workspace files.
    document_set = get_document_set_or_404(session, document_set_id)
    try:
        session.delete(document_set)
        session.commit()
    except Exception:
        session.rollback()
        raise

    for directory in (
        settings.data_dir / "originals" / document_set_id,
        settings.data_dir / "renders" / document_set_id,
        settings.data_dir / "generated" / document_set_id,
    ):
        shutil.rmtree(directory, ignore_errors=True)

    return {"deleted_id": document_set_id, "deleted": True}


def search_document_set(
    session: Session,
    document_set_id: str,
    query: str,
    limit: int = 50,
) -> dict:
    # Search extracted text across every current document in one set.
    get_document_set_or_404(session, document_set_id)
    cleaned_query = WHITESPACE_PATTERN.sub(" ", query).strip()
    if len(cleaned_query) < 2:
        return {
            "query": cleaned_query,
            "results": [],
            "result_count": 0,
            "truncated": False,
        }

    rows = session.execute(
        select(DocumentElement, DocumentRecord)
        .join(DocumentRecord, DocumentElement.document_id == DocumentRecord.id)
        .where(
            DocumentRecord.document_set_id == document_set_id,
            func.lower(DocumentElement.text).contains(cleaned_query.casefold()),
        )
        .order_by(
            DocumentRecord.original_name,
            DocumentElement.paragraph_index,
        )
        .limit(limit + 1)
    ).all()

    truncated = len(rows) > limit
    rows = rows[:limit]
    results = [
        {
            "element_id": element.id,
            "document_id": document.id,
            "document_name": document.original_name,
            "paragraph_index": element.paragraph_index,
            "element_type": element_type_for_style(element.style_name),
            **element_location_payload(element.style_name),
            "text": element.text,
        }
        for element, document in rows
    ]
    return {
        "query": cleaned_query,
        "results": results,
        "result_count": len(results),
        "truncated": truncated,
    }

def _create_exact_link_groups(session: Session, document_set_id: str) -> None:
    elements = session.scalars(
        select(DocumentElement)
        .join(DocumentRecord)
        .where(DocumentRecord.document_set_id == document_set_id)
    ).all()

    by_text: dict[tuple[str, str], list[DocumentElement]] = defaultdict(list)
    for element in elements:
        if element.normalized_text:
            by_text[
                (element.normalized_text, element_type_for_style(element.style_name))
            ].append(element)

    records: list[object] = []
    for (normalized_text, _element_type), matches in by_text.items():
        distinct_documents = {match.document_id for match in matches}
        if len(distinct_documents) < 2:
            continue

        group = LinkGroup(
            id=new_id(),
            document_set_id=document_set_id,
            representative_text=matches[0].text,
            normalized_text=normalized_text,
            match_type="exact",
        )
        records.append(group)
        records.extend(
            LinkMember(id=new_id(), link_group=group, element=element)
            for element in matches
        )

    if records:
        session.add_all(records)


def _rebuild_exact_link_groups(session: Session, document_set_id: str) -> None:
    """Refresh exact-match groups after the working document versions change."""

    groups = session.scalars(
        select(LinkGroup).where(LinkGroup.document_set_id == document_set_id)
    ).all()
    for group in groups:
        session.delete(group)
    session.flush()
    _create_exact_link_groups(session, document_set_id)

def list_document_sets(session: Session) -> dict:
    """Return lightweight saved-workspace summaries."""

    document_count = (
        select(func.count(DocumentRecord.id))
        .where(DocumentRecord.document_set_id == DocumentSet.id)
        .correlate(DocumentSet)
        .scalar_subquery()
    )

    edit_count = (
        select(func.count(GenerationJob.id))
        .where(
            GenerationJob.document_set_id == DocumentSet.id,
            GenerationJob.status == "completed",
        )
        .correlate(DocumentSet)
        .scalar_subquery()
    )

    rows = session.execute(
        select(
            DocumentSet.id.label("id"),
            DocumentSet.name.label("name"),
            DocumentSet.created_at.label("created_at"),
            document_count.label("document_count"),
            edit_count.label("edit_count"),
        ).order_by(DocumentSet.created_at.desc())
    ).mappings().all()

    return {
        "document_sets": [
            {
                "id": row["id"],
                "name": row["name"],
                "created_at": utc_isoformat(row["created_at"]),
                "document_count": int(row["document_count"] or 0),
                "edit_count": int(row["edit_count"] or 0),
            }
            for row in rows
        ]
    }
    
    
def get_document_set_or_404(session: Session, document_set_id: str) -> DocumentSet:
    group_count = int(
        session.scalar(
            select(func.count(LinkGroup.id)).where(
                LinkGroup.document_set_id == document_set_id
            )
        )
        or 0
    )
    element_count_rows = session.execute(
        select(DocumentElement.document_id, func.count(DocumentElement.id))
        .join(DocumentRecord, DocumentElement.document_id == DocumentRecord.id)
        .where(DocumentRecord.document_set_id == document_set_id)
        .group_by(DocumentElement.document_id)
    ).all()

    loader_options = [selectinload(DocumentSet.documents)]
    if group_count <= LARGE_SET_LINK_GROUP_LIMIT:
        loader_options.append(
            selectinload(DocumentSet.link_groups)
            .selectinload(LinkGroup.members)
            .selectinload(LinkMember.element)
            .selectinload(DocumentElement.document)
        )

    document_set = session.scalar(
        select(DocumentSet)
        .where(DocumentSet.id == document_set_id)
        .options(*loader_options)
    )
    if document_set is None:
        raise HTTPException(status_code=404, detail="Document set not found.")

    document_set._docsync_link_group_count = group_count
    document_set._docsync_element_counts = {
        document_id: int(element_count)
        for document_id, element_count in element_count_rows
    }
    return document_set


def serialize_document_set(document_set: DocumentSet) -> dict:
    documents = sorted(document_set.documents, key=lambda item: item.original_name.casefold())

    known_group_count = getattr(document_set, "_docsync_link_group_count", None)
    group_count = (
        int(known_group_count)
        if known_group_count is not None
        else len(document_set.link_groups)
    )
    groups = (
        sorted(
            document_set.link_groups,
            key=lambda item: (-len(item.members), item.representative_text.casefold()),
        )
        if group_count <= LARGE_SET_LINK_GROUP_LIMIT
        else []
    )

    element_counts = getattr(document_set, "_docsync_element_counts", None)

    def element_count(document: DocumentRecord) -> int:
        if element_counts is not None:
            return int(element_counts.get(document.id, 0))
        return len(document.elements)

    return {
        "id": document_set.id,
        "name": document_set.name,
        "created_at": utc_isoformat(document_set.created_at),
        "documents": [
            {
                "id": document.id,
                "version_id": document.id,
                "name": document.original_name,
                "checksum_sha256": document.checksum_sha256,
                "element_count": element_count(document),
                "view_url": f"/api/document-versions/{document.id}/pages",
            }
            for document in documents
        ],
        "link_group_count": group_count,
        "link_groups": [serialize_link_group(group) for group in groups],
    }


def get_document_or_404(session: Session, document_id: str) -> DocumentRecord:
    document = session.scalar(
        select(DocumentRecord)
        .where(DocumentRecord.id == document_id)
        .options(selectinload(DocumentRecord.elements))
    )
    if document is None:
        raise HTTPException(status_code=404, detail="Document version not found.")
    return document


def serialize_document_view(document: DocumentRecord) -> dict:
    pages = _layout_pages(document.elements)
    return {
        "document_id": document.id,
        "version_id": document.id,
        "document_set_id": document.document_set_id,
        "document_name": document.original_name,
        "render_status": "ready",
        "render_mode": "structured",
        "pagination": "estimated",
        "page_count": len(pages),
        "notice": (
            "Structured browser preview. Paragraphs, headings, list items, and non-empty "
            "top-level table cells are selectable. Page grouping is estimated; the original "
            "DOCX remains the source of truth."
        ),
        "pages": pages,
    }


def rendered_pdf_path(document: DocumentRecord) -> Path:
    return settings.data_dir / "renders" / document.document_set_id / f"{document.id}.pdf"


def current_document_path(session: Session, document: DocumentRecord) -> Path:
    """Return the latest applied version, falling back to the immutable upload."""

    latest_version = session.scalar(
        select(GeneratedVersion)
        .join(GenerationJob, GeneratedVersion.generation_id == GenerationJob.id)
        .where(
            GeneratedVersion.source_document_id == document.id,
            GenerationJob.status == "completed",
        )
        .order_by(GenerationJob.created_at.desc(), GeneratedVersion.id.desc())
        .limit(1)
    )
    if latest_version is not None:
        path = settings.data_dir / "generated" / latest_version.storage_name
    else:
        path = settings.data_dir / "originals" / document.stored_name
    if not path.exists():
        raise HTTPException(status_code=500, detail="The current document file is missing.")
    return path


def render_document_with_word(session: Session, document: DocumentRecord) -> dict:
    """Render the current DOCX version through Microsoft Word and cache the PDF."""

    source_path = current_document_path(session, document)

    output_path = rendered_pdf_path(document)
    if not output_path.exists() or output_path.stat().st_size == 0:
        powershell = shutil.which("powershell.exe") or shutil.which("powershell")
        render_script = settings.render_script
        if powershell is None or not render_script.exists():
            raise HTTPException(
                status_code=503,
                detail="Microsoft Word rendering is unavailable on this server.",
            )

        output_path.parent.mkdir(parents=True, exist_ok=True)
        temporary_path = output_path.with_name(f"{document.id}-{new_id()}.tmp.pdf")
        try:
            with WORD_RENDER_LOCK:
                if output_path.exists() and output_path.stat().st_size > 0:
                    temporary_path.unlink(missing_ok=True)
                else:
                    result = subprocess.run(
                        [
                            powershell,
                            "-NoProfile",
                            "-NonInteractive",
                            "-ExecutionPolicy",
                            "Bypass",
                            "-File",
                            str(render_script),
                            "-SourcePath",
                            str(source_path),
                            "-OutputPath",
                            str(temporary_path),
                        ],
                        capture_output=True,
                        text=True,
                        timeout=120,
                        check=False,
                    )
                    if result.returncode != 0 or not temporary_path.exists():
                        raise HTTPException(
                            status_code=422,
                            detail=(
                                "Microsoft Word could not render this document. Close any Word "
                                "dialogues and confirm the file opens normally in Word."
                            ),
                        )
                    temporary_path.replace(output_path)
        except subprocess.TimeoutExpired as exc:
            raise HTTPException(
                status_code=504,
                detail="Microsoft Word took too long to render this document.",
            ) from exc
        finally:
            temporary_path.unlink(missing_ok=True)

    view = serialize_document_view(document)
    view.update(
        {
            "render_mode": "word_pdf",
            "pagination": "word",
            "pdf_url": (
                f"/api/document-versions/{document.id}/rendered-file"
                f"?v={output_path.stat().st_mtime_ns}"
            ),
            "notice": (
                "Word layout preview generated by Microsoft Word. Use Select text mode to "
                "choose a supported element for synchronisation."
            ),
        }
    )
    return view


def serialize_link_group(group: LinkGroup) -> dict:
    members = sorted(
        group.members,
        key=lambda member: (
            member.element.document.original_name.casefold(),
            member.element.paragraph_index,
        ),
    )
    return {
        "id": group.id,
        "match_type": group.match_type,
        "representative_text": group.representative_text,
        "member_count": len(members),
        "document_count": len({member.element.document_id for member in members}),
        "members": [
            {
                "element_id": member.element.id,
                "document_id": member.element.document_id,
                "document_name": member.element.document.original_name,
                "paragraph_index": member.element.paragraph_index,
                "element_type": element_type_for_style(member.element.style_name),
                "text": member.element.text,
                "style_name": display_style_name(member.element.style_name),
                **element_location_payload(member.element.style_name),
            }
            for member in members
        ],
    }


def get_link_group_or_404(
    session: Session,
    document_set_id: str,
    link_group_id: str,
) -> LinkGroup:
    group = session.scalar(
        select(LinkGroup)
        .where(
            LinkGroup.id == link_group_id,
            LinkGroup.document_set_id == document_set_id,
        )
        .options(
            selectinload(LinkGroup.members)
            .selectinload(LinkMember.element)
            .selectinload(DocumentElement.document)
        )
    )
    if group is None:
        raise HTTPException(status_code=404, detail="Link group not found in this document set.")
    return group


def get_element_matches_or_404(session: Session, element_id: str) -> dict:
    element = session.scalar(
        select(DocumentElement)
        .where(DocumentElement.id == element_id)
        .options(selectinload(DocumentElement.document))
    )
    if element is None:
        raise HTTPException(status_code=404, detail="Document element not found.")

    group = session.scalar(
        select(LinkGroup)
        .join(LinkMember)
        .where(LinkMember.element_id == element_id)
        .options(
            selectinload(LinkGroup.members)
            .selectinload(LinkMember.element)
            .selectinload(DocumentElement.document)
        )
    )
    source = {
        "element_id": element.id,
        "document_id": element.document_id,
        "document_name": element.document.original_name,
        "paragraph_index": element.paragraph_index,
        "element_type": element_type_for_style(element.style_name),
        "text": element.text,
        "style_name": display_style_name(element.style_name),
        **element_location_payload(element.style_name),
    }
    return {
        "source": source,
        "link_group": serialize_link_group(group) if group is not None else None,
        "exact_match_count": len(group.members) - 1 if group is not None else 0,
        "similar_matches": [],
        "similarity_status": "not_enabled",
    }


def _selected_group_elements(
    group: LinkGroup,
    included_element_ids: list[str] | None,
    source_element_id: str | None = None,
) -> list[DocumentElement]:
    members_by_id = {member.element.id: member.element for member in group.members}
    if included_element_ids is None:
        selected = list(members_by_id.values())
    else:
        unknown = set(included_element_ids) - set(members_by_id)
        if unknown:
            raise HTTPException(
                status_code=422,
                detail="One or more included elements are not members of this exact-match group.",
            )
        selected = [members_by_id[element_id] for element_id in included_element_ids]

    if source_element_id is not None:
        if source_element_id not in members_by_id:
            raise HTTPException(
                status_code=422,
                detail="The source element is not a member of this exact-match group.",
            )
        if source_element_id not in {element.id for element in selected}:
            raise HTTPException(
                status_code=422,
                detail="The source element must remain included in the change.",
            )
    if not selected:
        raise HTTPException(status_code=422, detail="Include at least one target element.")
    return selected


def preview_edit(
    group: LinkGroup,
    replacement_text: str,
    included_element_ids: list[str] | None = None,
    source_element_id: str | None = None,
) -> dict:
    by_document: dict[str, dict] = {}
    elements = _selected_group_elements(group, included_element_ids, source_element_id)
    for element in elements:
        item = by_document.setdefault(
            element.document_id,
            {
                "document_id": element.document_id,
                "document_name": element.document.original_name,
                "changes": [],
            },
        )
        item["changes"].append(
            {
                "element_id": element.id,
                "paragraph_index": element.paragraph_index,
                "element_type": element_type_for_style(element.style_name),
                **element_location_payload(element.style_name),
                "before": element.text,
                "after": replacement_text,
            }
        )

    documents = sorted(by_document.values(), key=lambda item: item["document_name"].casefold())
    for document in documents:
        document["changes"].sort(key=lambda item: item["paragraph_index"])

    return {
        "link_group_id": group.id,
        "source_element_id": source_element_id,
        "replacement_text": replacement_text,
        "affected_document_count": len(documents),
        "affected_location_count": sum(len(document["changes"]) for document in documents),
        "documents": documents,
    }


def _replace_paragraph_text(paragraph: Paragraph, replacement_text: str) -> None:
    # Keep paragraph-level styling and reuse the first run's formatting as the MVP policy.
    if paragraph.runs:
        paragraph.runs[0].text = replacement_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(replacement_text)


def _replace_table_cell_text(cell, replacement_text: str) -> None:
    paragraphs = list(cell.paragraphs)
    if not paragraphs:
        cell.text = replacement_text
        return

    _replace_paragraph_text(paragraphs[0], replacement_text)
    for paragraph in paragraphs[1:]:
        _replace_paragraph_text(paragraph, "")


def _load_source_document(session: Session, record: DocumentRecord) -> DocxDocument:
    return Document(current_document_path(session, record))


def generate_versions(
    session: Session,
    document_set_id: str,
    group: LinkGroup,
    replacement_text: str,
    included_element_ids: list[str] | None = None,
    source_element_id: str | None = None,
) -> GenerationJob:
    generation_id = new_id()
    generation_dir = settings.data_dir / "generated" / document_set_id / generation_id
    generation_dir.mkdir(parents=True, exist_ok=False)

    selected_elements = _selected_group_elements(
        group, included_element_ids, source_element_id
    )
    members_by_document: dict[str, list[DocumentElement]] = defaultdict(list)
    for element in selected_elements:
        members_by_document[element.document_id].append(element)

    generated_files: list[tuple[DocumentRecord, str, Path]] = []
    try:
        for document_id, elements in members_by_document.items():
            record = elements[0].document
            docx = _load_source_document(session, record)
            for element in sorted(elements, key=lambda item: item.paragraph_index):
                table_location = table_cell_location(element.style_name)
                if table_location is not None:
                    table_index, row_index, column_index = table_location
                    try:
                        cell = docx.tables[table_index].rows[row_index].cells[column_index]
                    except IndexError as exc:
                        raise HTTPException(
                            status_code=500,
                            detail=(
                                f"Table-cell location is no longer valid for "
                                f"{record.original_name}."
                            ),
                        ) from exc
                    _replace_table_cell_text(cell, replacement_text)
                    continue

                if element.paragraph_index >= len(docx.paragraphs):
                    raise HTTPException(
                        status_code=500,
                        detail=f"Paragraph location is no longer valid for {record.original_name}.",
                    )
                _replace_paragraph_text(docx.paragraphs[element.paragraph_index], replacement_text)

            source_name = Path(record.original_name)
            output_name = safe_download_name(f"{source_name.stem}-updated.docx")
            output_path = generation_dir / output_name
            docx.save(output_path)
            generated_files.append((record, output_name, output_path))

        # Keep the stored name short for Windows workspaces that are already near MAX_PATH.
        zip_name = "current-documents.zip"
        job = GenerationJob(
            id=generation_id,
            document_set_id=document_set_id,
            link_group_id=group.id,
            replacement_text=replacement_text,
            zip_storage_name=f"{document_set_id}/{generation_id}/{zip_name}",
            status="completed",
        )
        session.add(job)
        for element in selected_elements:
            session.add(
                GenerationTarget(
                    id=new_id(),
                    generation=job,
                    element_id=element.id,
                    document_id=element.document_id,
                    document_name=element.document.original_name,
                    paragraph_index=element.paragraph_index,
                    before_text=element.text,
                    after_text=replacement_text,
                )
            )
        for record, output_name, output_path in generated_files:
            session.add(
                GeneratedVersion(
                    id=new_id(),
                    generation=job,
                    source_document_id=record.id,
                    download_name=output_name,
                    storage_name=str(output_path.relative_to(settings.data_dir / "generated")),
                )
            )
        session.flush()

        # The in-workspace element model now follows the applied DOCX versions. Stable
        # document IDs are retained, while exact-match groups are rebuilt from current text.
        for element in selected_elements:
            element.text = replacement_text
            element.normalized_text = normalise_text(replacement_text)
        _rebuild_exact_link_groups(session, document_set_id)

        changed_paths = {record.id: path for record, _, path in generated_files}
        documents = session.scalars(
            select(DocumentRecord)
            .where(DocumentRecord.document_set_id == document_set_id)
            .order_by(DocumentRecord.original_name)
        ).all()
        zip_path = generation_dir / zip_name
        with zipfile.ZipFile(zip_path, mode="w", compression=zipfile.ZIP_DEFLATED) as archive:
            for document in documents:
                snapshot_path = changed_paths.get(document.id)
                if snapshot_path is None:
                    snapshot_path = current_document_path(session, document)
                archive.write(snapshot_path, arcname=safe_download_name(document.original_name))

        for record, _, _ in generated_files:
            rendered_pdf_path(record).unlink(missing_ok=True)

        session.commit()
        session.refresh(job)
        return session.scalar(
            select(GenerationJob)
            .where(GenerationJob.id == generation_id)
            .options(
                selectinload(GenerationJob.versions),
                selectinload(GenerationJob.targets),
            )
        )
    except Exception:
        session.rollback()
        shutil.rmtree(generation_dir, ignore_errors=True)
        raise


def get_generation_or_404(session: Session, generation_id: str) -> GenerationJob:
    job = session.scalar(
        select(GenerationJob)
        .where(GenerationJob.id == generation_id)
        .options(
            selectinload(GenerationJob.versions),
            selectinload(GenerationJob.targets),
        )
    )
    if job is None:
        raise HTTPException(status_code=404, detail="Generation not found.")
    return job


def serialize_document_set_history(session: Session, document_set_id: str) -> dict:
    get_document_set_or_404(session, document_set_id)
    jobs = session.scalars(
        select(GenerationJob)
        .where(GenerationJob.document_set_id == document_set_id)
        .order_by(GenerationJob.created_at.desc())
        .options(
            selectinload(GenerationJob.versions),
            selectinload(GenerationJob.targets),
        )
    ).all()
    return {
        "document_set_id": document_set_id,
        "events": [
            {
                "generation_id": job.id,
                "event_type": "synchronised_edit",
                "status": job.status,
                "created_at": utc_isoformat(job.created_at),
                "replacement_text": job.replacement_text,
                "target_count": len(job.targets),
                "version_count": len(job.versions),
                "targets": [
                    {
                        "element_id": target.element_id,
                        "document_id": target.document_id,
                        "document_name": target.document_name,
                        "paragraph_index": target.paragraph_index,
                        "before": target.before_text,
                        "after": target.after_text,
                    }
                    for target in sorted(
                        job.targets,
                        key=lambda item: (
                            item.document_name.casefold(),
                            item.paragraph_index,
                        ),
                    )
                ],
            }
            for job in jobs
        ],
    }


def generation_download_path(job: GenerationJob) -> Path:
    path = settings.data_dir / "generated" / job.zip_storage_name
    if not path.exists():
        raise HTTPException(status_code=404, detail="Generated ZIP file is missing.")
    return path
