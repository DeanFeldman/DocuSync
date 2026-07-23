from __future__ import annotations

import hashlib
import importlib
import io
import os
import sys
import zipfile
from pathlib import Path

API_DIR = Path(__file__).resolve().parents[1]
if str(API_DIR) not in sys.path:
    sys.path.insert(0, str(API_DIR))

from docx import Document
from fastapi.testclient import TestClient


def make_docx(building: str, address: str, shared: str) -> bytes:
    stream = io.BytesIO()
    document = Document()
    document.add_heading(f"{building} Agreement", level=1)
    document.add_paragraph(f"Property: {building}")
    document.add_paragraph(f"Address: {address}")
    document.add_paragraph(shared)
    document.add_paragraph(f"Unique contact for {building}.")
    document.save(stream)
    return stream.getvalue()


def make_typed_docx(title: str, shared: str, shared_is_heading: bool) -> bytes:
    stream = io.BytesIO()
    document = Document()
    document.add_heading(title, level=1)
    if shared_is_heading:
        document.add_heading(shared, level=2)
    else:
        document.add_paragraph(shared)
    document.save(stream)
    return stream.getvalue()


def load_test_app(tmp_path: Path):
    os.environ["DOCUMENTSYNC_DATA_DIR"] = str(tmp_path / "data")
    os.environ["DOCUMENTSYNC_DATABASE_URL"] = f"sqlite:///{tmp_path / 'test.db'}"

    for module_name in list(sys.modules):
        if module_name == "app" or module_name.startswith("app."):
            del sys.modules[module_name]

    main = importlib.import_module("app.main")
    return main.app


def test_desktop_session_cookie_protects_document_endpoints(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setenv("DOCUMENTSYNC_SESSION_TOKEN", "desktop-test-token")
    app = load_test_app(tmp_path)

    with TestClient(app) as client:
        assert client.get("/api/health").status_code == 200
        denied = client.post("/api/document-sets")
        assert denied.status_code == 401
        assert denied.json()["detail"] == "The desktop session is missing or invalid."

        client.cookies.set("docsync_session", "desktop-test-token")
        accepted = client.post("/api/document-sets")
        assert accepted.status_code == 422


def test_built_desktop_frontend_is_served_by_the_api(tmp_path: Path, monkeypatch) -> None:
    web_dist = tmp_path / "web"
    (web_dist / "assets").mkdir(parents=True)
    (web_dist / "index.html").write_text("<h1>DocSync desktop</h1>", encoding="utf-8")
    (web_dist / "assets" / "app.js").write_text("console.log('ready')", encoding="utf-8")
    monkeypatch.setenv("DOCUMENTSYNC_WEB_DIST", str(web_dist))
    monkeypatch.delenv("DOCUMENTSYNC_SESSION_TOKEN", raising=False)
    app = load_test_app(tmp_path)

    with TestClient(app) as client:
        page = client.get("/")
        assert page.status_code == 200
        assert "DocSync desktop" in page.text
        assert page.headers["x-content-type-options"] == "nosniff"
        assert client.get("/assets/app.js").status_code == 200


def test_saved_document_set_library_lists_and_reopens_workspaces(tmp_path: Path) -> None:
    shared = "The building manager must submit the report every month."
    app = load_test_app(tmp_path)

    with TestClient(app) as client:
        empty_library = client.get("/api/document-sets")
        assert empty_library.status_code == 200
        assert empty_library.json() == {"document_sets": []}

        uploaded = client.post(
            "/api/document-sets",
            data={"name": "Saved building agreements"},
            files=[
                (
                    "files",
                    (
                        "Building-A.docx",
                        io.BytesIO(make_docx("Building A", "1 Alpha Road", shared)),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
                (
                    "files",
                    (
                        "Building-B.docx",
                        io.BytesIO(make_docx("Building B", "2 Beta Avenue", shared)),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
            ],
        )
        assert uploaded.status_code == 201, uploaded.text
        workspace = uploaded.json()
        group = next(
            item for item in workspace["link_groups"] if item["representative_text"] == shared
        )

        generated = client.post(
            f"/api/document-sets/{workspace['id']}/generate",
            json={
                "link_group_id": group["id"],
                "replacement_text": "The building manager must submit the report every Friday.",
            },
        )
        assert generated.status_code == 201, generated.text

        library = client.get("/api/document-sets")
        assert library.status_code == 200
        assert library.json()["document_sets"] == [
            {
                "id": workspace["id"],
                "name": "Saved building agreements",
                "created_at": workspace["created_at"],
                "document_count": 2,
                "edit_count": 1,
            }
        ]

        reopened = client.get(f"/api/document-sets/{workspace['id']}")
        assert reopened.status_code == 200
        assert reopened.json()["id"] == workspace["id"]
        assert len(reopened.json()["documents"]) == 2


def test_exact_match_preview_generation_and_original_immutability(tmp_path: Path) -> None:
    shared = "The building manager must submit the report every month."
    replacement = (
        "The building manager must submit the report by the fifth working day of every month."
    )
    originals = {
        "Building-A-Agreement.docx": make_docx("Building A", "1 Alpha Road", shared),
        "Building-B-Agreement.docx": make_docx("Building B", "2 Beta Avenue", shared),
        "Building-C-Agreement.docx": make_docx("Building C", "3 Gamma Street", shared),
    }
    original_hashes = {name: hashlib.sha256(payload).hexdigest() for name, payload in originals.items()}

    app = load_test_app(tmp_path)
    with TestClient(app) as client:
        preflight = client.options(
            "/api/document-sets",
            headers={
                "Origin": "http://localhost:5174",
                "Access-Control-Request-Method": "POST",
            },
        )
        assert preflight.status_code == 200
        assert preflight.headers["access-control-allow-origin"] == "http://localhost:5174"

        response = client.post(
            "/api/document-sets",
            data={"name": "Building agreements"},
            files=[
                (
                    "files",
                    (name, io.BytesIO(payload), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                )
                for name, payload in originals.items()
            ],
        )
        assert response.status_code == 201, response.text
        payload = response.json()
        assert len(payload["documents"]) == 3

        matching_group = next(
            group for group in payload["link_groups"] if group["representative_text"] == shared
        )
        assert matching_group["document_count"] == 3

        document_view = client.get(payload["documents"][0]["view_url"])
        assert document_view.status_code == 200
        assert document_view.json()["render_status"] == "ready"
        assert document_view.json()["pagination"] == "estimated"
        assert any(
            element["id"]
            for page in document_view.json()["pages"]
            for element in page["elements"]
        )

        config = importlib.import_module("app.config")
        rendered_path = (
            config.settings.data_dir
            / "renders"
            / payload["id"]
            / f"{payload['documents'][0]['version_id']}.pdf"
        )
        rendered_path.parent.mkdir(parents=True, exist_ok=True)
        rendered_path.write_bytes(b"%PDF-1.4\n% DocumentSync cached render fixture\n")
        rendered = client.post(
            f"/api/documents/{payload['documents'][0]['version_id']}/render"
        )
        assert rendered.status_code == 200, rendered.text
        assert rendered.json()["render_mode"] == "word_pdf"
        assert rendered.json()["pagination"] == "word"
        pdf_response = client.get(rendered.json()["pdf_url"])
        assert pdf_response.status_code == 200
        assert pdf_response.headers["content-type"] == "application/pdf"

        source_element_id = matching_group["members"][0]["element_id"]
        matches = client.get(f"/api/document-elements/{source_element_id}/matches")
        assert matches.status_code == 200
        assert matches.json()["exact_match_count"] == 2
        assert matches.json()["link_group"]["id"] == matching_group["id"]

        preview = client.post(
            f"/api/document-sets/{payload['id']}/preview",
            json={
                "link_group_id": matching_group["id"],
                "replacement_text": replacement,
            },
        )
        assert preview.status_code == 200
        assert preview.json()["affected_document_count"] == 3
        assert preview.json()["affected_location_count"] == 3

        generated = client.post(
            f"/api/document-sets/{payload['id']}/generate",
            json={
                "link_group_id": matching_group["id"],
                "replacement_text": replacement,
            },
        )
        assert generated.status_code == 201, generated.text
        updated_set = generated.json()["document_set"]
        refreshed_group = next(
            group
            for group in updated_set["link_groups"]
            if group["representative_text"] == replacement
        )
        refreshed_view = client.get(updated_set["documents"][0]["view_url"])
        assert refreshed_view.status_code == 200
        assert any(
            element["text"] == replacement
            for page in refreshed_view.json()["pages"]
            for element in page["elements"]
        )

        current_docx = client.get(
            f"/api/documents/{updated_set['documents'][0]['id']}/download"
        )
        assert current_docx.status_code == 200
        current_document = Document(io.BytesIO(current_docx.content))
        assert replacement in "\n".join(
            paragraph.text for paragraph in current_document.paragraphs
        )

        second_replacement = "The building manager must submit the report every Monday."
        second_generation = client.post(
            f"/api/document-sets/{payload['id']}/generate",
            json={
                "link_group_id": refreshed_group["id"],
                "replacement_text": second_replacement,
            },
        )
        assert second_generation.status_code == 201, second_generation.text
        assert any(
            group["representative_text"] == second_replacement
            for group in second_generation.json()["document_set"]["link_groups"]
        )
        download = client.get(second_generation.json()["download_url"])
        assert download.status_code == 200

        history = client.get(f"/api/document-sets/{payload['id']}/history")
        assert history.status_code == 200
        assert len(history.json()["events"]) == 2
        assert history.json()["events"][0]["target_count"] == 3
        assert history.json()["events"][0]["version_count"] == 3

    with zipfile.ZipFile(io.BytesIO(download.content)) as archive:
        assert len(archive.namelist()) == 3
        for member_name in archive.namelist():
            updated = Document(io.BytesIO(archive.read(member_name)))
            all_text = "\n".join(paragraph.text for paragraph in updated.paragraphs)
            assert second_replacement in all_text
            building_letter = member_name.split("-")[1]
            assert f"Building {building_letter}" in all_text

    originals_dir = tmp_path / "data" / "originals" / payload["id"]
    stored_files = list(originals_dir.glob("*.docx"))
    assert len(stored_files) == 3
    stored_hashes = {hashlib.sha256(path.read_bytes()).hexdigest() for path in stored_files}
    assert stored_hashes == set(original_hashes.values())


def test_explicit_target_confirmation_controls_preview_and_generation(tmp_path: Path) -> None:
    shared = "The fire register must be reviewed every Friday."
    replacement = "The fire register must be reviewed before noon every Friday."
    originals = {
        "Building-A.docx": make_docx("Building A", "1 Alpha Road", shared),
        "Building-B.docx": make_docx("Building B", "2 Beta Avenue", shared),
        "Building-C.docx": make_docx("Building C", "3 Gamma Street", shared),
    }

    app = load_test_app(tmp_path)
    with TestClient(app) as client:
        uploaded = client.post(
            "/api/document-sets",
            data={"name": "Safety agreements"},
            files=[
                (
                    "files",
                    (
                        name,
                        io.BytesIO(payload),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                )
                for name, payload in originals.items()
            ],
        )
        assert uploaded.status_code == 201, uploaded.text
        document_set = uploaded.json()
        group = next(
            item for item in document_set["link_groups"] if item["representative_text"] == shared
        )
        members_by_name = {member["document_name"]: member for member in group["members"]}
        source_id = members_by_name["Building-A.docx"]["element_id"]
        included_ids = [
            source_id,
            members_by_name["Building-B.docx"]["element_id"],
        ]
        request = {
            "link_group_id": group["id"],
            "source_element_id": source_id,
            "included_element_ids": included_ids,
            "replacement_text": replacement,
        }

        preview = client.post(
            f"/api/document-sets/{document_set['id']}/preview",
            json=request,
        )
        assert preview.status_code == 200, preview.text
        assert preview.json()["affected_document_count"] == 2
        assert {item["document_name"] for item in preview.json()["documents"]} == {
            "Building-A.docx",
            "Building-B.docx",
        }

        source_excluded = client.post(
            f"/api/document-sets/{document_set['id']}/preview",
            json={**request, "included_element_ids": included_ids[1:]},
        )
        assert source_excluded.status_code == 422

        outsider = client.post(
            f"/api/document-sets/{document_set['id']}/preview",
            json={**request, "included_element_ids": [source_id, "not-a-group-member"]},
        )
        assert outsider.status_code == 422

        generated = client.post(
            f"/api/document-sets/{document_set['id']}/generate",
            json=request,
        )
        assert generated.status_code == 201, generated.text
        assert {item["name"] for item in generated.json()["files"]} == {
            "Building-A-updated.docx",
            "Building-B-updated.docx",
        }

        history = client.get(f"/api/document-sets/{document_set['id']}/history")
        assert history.status_code == 200
        assert history.json()["events"][0]["target_count"] == 2
        assert {item["document_name"] for item in history.json()["events"][0]["targets"]} == {
            "Building-A.docx",
            "Building-B.docx",
        }


def test_exact_text_does_not_link_incompatible_element_types(tmp_path: Path) -> None:
    shared = "Quarterly reporting"
    app = load_test_app(tmp_path)
    with TestClient(app) as client:
        uploaded = client.post(
            "/api/document-sets",
            data={"name": "Mixed structures"},
            files=[
                (
                    "files",
                    (
                        "Heading.docx",
                        io.BytesIO(make_typed_docx("Heading source", shared, True)),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
                (
                    "files",
                    (
                        "Paragraph.docx",
                        io.BytesIO(make_typed_docx("Paragraph source", shared, False)),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
            ],
        )
        assert uploaded.status_code == 201, uploaded.text
        assert all(
            group["representative_text"] != shared for group in uploaded.json()["link_groups"]
        )

def test_set_management_and_global_search(tmp_path: Path) -> None:
    shared = "The manager must submit a monthly compliance report."
    app = load_test_app(tmp_path)

    with TestClient(app) as client:
        uploaded = client.post(
            "/api/document-sets",
            data={"name": "Managed agreements"},
            files=[
                (
                    "files",
                    (
                        "Alpha.docx",
                        io.BytesIO(make_docx("Alpha", "1 Alpha Road", shared)),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
                (
                    "files",
                    (
                        "Beta.docx",
                        io.BytesIO(make_docx("Beta", "2 Beta Road", shared)),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
            ],
        )
        assert uploaded.status_code == 201, uploaded.text
        workspace = uploaded.json()

        added = client.post(
            f"/api/document-sets/{workspace['id']}/documents",
            files=[
                (
                    "files",
                    (
                        "Gamma.docx",
                        io.BytesIO(
                            make_docx(
                                "Gamma",
                                "3 Gamma Road",
                                "Gamma contains the searchable emergency procedure.",
                            )
                        ),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                )
            ],
        )
        assert added.status_code == 201, added.text
        managed = added.json()
        assert len(managed["documents"]) == 3
        gamma = next(
            document for document in managed["documents"]
            if document["name"] == "Gamma.docx"
        )

        search = client.get(
            f"/api/document-sets/{workspace['id']}/search",
            params={"q": "searchable emergency"},
        )
        assert search.status_code == 200, search.text
        assert search.json()["result_count"] == 1
        assert search.json()["results"][0]["document_id"] == gamma["id"]

        removed = client.delete(
            f"/api/document-sets/{workspace['id']}/documents/{gamma['id']}"
        )
        assert removed.status_code == 200, removed.text
        assert len(removed.json()["documents"]) == 2

        blocked = client.delete(
            f"/api/document-sets/{workspace['id']}/documents/"
            f"{removed.json()['documents'][0]['id']}"
        )
        assert blocked.status_code == 422
        assert "at least two documents" in blocked.json()["detail"]

        no_results = client.get(
            f"/api/document-sets/{workspace['id']}/search",
            params={"q": "searchable emergency"},
        )
        assert no_results.status_code == 200
        assert no_results.json()["results"] == []

        deleted = client.delete(f"/api/document-sets/{workspace['id']}")
        assert deleted.status_code == 200
        assert deleted.json() == {
            "deleted_id": workspace["id"],
            "deleted": True,
        }
        assert client.get(f"/api/document-sets/{workspace['id']}").status_code == 404
        assert client.get("/api/document-sets").json() == {"document_sets": []}
        assert not (tmp_path / "data" / "originals" / workspace["id"]).exists()
