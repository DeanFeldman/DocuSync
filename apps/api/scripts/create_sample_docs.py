from pathlib import Path

from docx import Document


OUTPUT_DIR = Path(__file__).resolve().parents[1] / "sample-documents"
SHARED = "The building manager must submit the report every month."


def create_agreement(name: str, address: str) -> None:
    document = Document()
    document.add_heading(f"{name} Agreement", level=1)
    document.add_paragraph(f"Property: {name}")
    document.add_paragraph(f"Address: {address}")
    document.add_paragraph(SHARED)
    document.add_paragraph(f"The emergency contact for {name} is unique to this agreement.")
    document.save(OUTPUT_DIR / f"{name.replace(' ', '-')}-Agreement.docx")


if __name__ == "__main__":
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    create_agreement("Building A", "1 Alpha Road")
    create_agreement("Building B", "2 Beta Avenue")
    create_agreement("Building C", "3 Gamma Street")
    print(f"Created sample documents in {OUTPUT_DIR}")
